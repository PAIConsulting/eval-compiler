"""
compile_report.py

Reads a CSV of evaluation responses and produces one Word report per domain.

Project:    Evaluation Compiler (InfoShare evaluation reporting tool)
Owner:      PAI Consulting
Author:     Amanda Rummel (with code generation assistance from Claude / Anthropic AI)
Version:    0.1.0 (demonstration)
Created:    2026
License:    Proprietary - see LICENSE file

================================================================================
AUDIT-RELEVANT NOTES
================================================================================

This script:
  - Runs entirely on the local machine
  - Reads ONE input file: a CSV of evaluation responses (path hardcoded to
    fake_responses.csv in the script's working directory)
  - Reads ONE configuration file: keywords.yaml (in the same directory)
  - Writes Word documents to the ./reports/ subdirectory
  - Does NOT make any network calls (no imports of requests, urllib, socket,
    httpx, http.client, etc.)
  - Does NOT use any AI, machine-learning, or LLM libraries (no imports of
    openai, anthropic, transformers, torch, tensorflow, langchain, etc.)
  - Does NOT execute external commands or shell processes
  - Does NOT use eval(), exec(), or other dynamic code execution
  - Uses yaml.safe_load() (not yaml.load()) for safe YAML parsing
  - Categorization is purely deterministic: rule-based string matching
    against a hand-maintained keyword dictionary

Dependencies (see requirements.txt for exact pinned versions):
  - python-docx: Word document generation
  - PyYAML:      YAML configuration file parsing
  - matplotlib:  Chart generation

================================================================================
USAGE
================================================================================

    python compile_report.py

Reads fake_responses.csv and keywords.yaml from the current directory.
Writes one Word document per domain to reports/<domain>_report.docx.
"""

import csv
import os
import yaml
from collections import Counter, defaultdict
from io import BytesIO

import matplotlib
matplotlib.use("Agg")  # No display needed
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ==================== CONFIGURATION ====================

INPUT_CSV = "fake_responses.csv"
KEYWORDS_FILE = "keywords.yaml"
OUTPUT_DIR = "reports"

# Brand-style colors (placeholder - replace with conference branding)
COLOR_HEADER = RGBColor(0x1F, 0x3A, 0x5F)  # navy
COLOR_ACCENT = RGBColor(0x2E, 0x75, 0xB6)  # blue
COLOR_LIGHT_BG = "DCE6F1"  # light blue (hex string for shading)
COLOR_DARKER_BG = "4F81BD"  # darker blue for subject headers

# ==================== DATA LOADING ====================

def load_responses(csv_path):
    """Load all responses from CSV."""
    with open(csv_path, "r", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def load_keywords(yaml_path):
    """Load keyword dictionary."""
    with open(yaml_path, "r", encoding="utf-8") as f:
        data = yaml.safe_load(f)
    return data.get("categories", [])


# ==================== KEYWORD CATEGORIZATION ====================

def categorize_comment(comment, categories):
    """
    Apply keyword dictionary to a single free-text comment.
    Returns (keyword_1, keyword_2) or ("Uncategorized", "—") if no match.
    
    Strategy: count how many triggers from each category appear; pick the
    category with the most matches. Ties broken by earliest-defined category.
    """
    if not comment or not comment.strip():
        return None  # empty comments don't get categorized at all
    
    text = comment.lower()
    best_match = None
    best_score = 0
    
    for cat in categories:
        score = sum(1 for trigger in cat["triggers"] if trigger.lower() in text)
        if score > best_score:
            best_score = score
            best_match = cat
    
    if best_match is None:
        return ("Uncategorized", "—")
    return (best_match["keyword_1"], best_match["keyword_2"])


# ==================== ANALYSIS PER DOMAIN ====================

def analyze_domain(responses, domain, categories):
    """
    Build a structured analysis for one domain. Returns a dict with:
      - respondents: count
      - first_time: count
      - q2_counts: Counter of topic -> votes
      - q3 / q4 / q5: each with subject_counts, topic_counts (by subject),
                      and categorized comments
      - q6 / q7 / q8: list of (comment, kw1, kw2) tuples, sorted
    """
    domain_rows = [r for r in responses if r["Q1_Domain"] == domain]
    
    out = {
        "domain": domain,
        "respondents": len(domain_rows),
        "first_time": sum(1 for r in domain_rows if r["Q9_FirstTime"] == "Yes"),
    }
    
    # Q2: multi-select vote count
    q2_counts = Counter()
    for r in domain_rows:
        topics = [t.strip() for t in r["Q2_Topics"].split(";") if t.strip()]
        q2_counts.update(topics)
    out["q2_counts"] = q2_counts
    
    # Q3, Q4, Q5: cascading subject + topic + comments
    for q in ("Q3", "Q4", "Q5"):
        subj_col = f"{q}_Subject"
        subj_other_col = f"{q}_Subject_Other"
        topic_col = f"{q}a_Topic"
        topic_other_col = f"{q}a_Topic_Other"
        comment_col = f"{q}b_Comments"
        
        subject_counts = Counter()
        topics_by_subject = defaultdict(Counter)
        comments = []  # list of (subject, topic, comment, kw1, kw2)
        
        for r in domain_rows:
            subj = r[subj_col]
            if subj == "Other" and r[subj_other_col]:
                subj_display = f"Other: {r[subj_other_col]}"
            else:
                subj_display = subj
            
            topic = r[topic_col]
            if topic == "Other" and r[topic_other_col]:
                topic_display = f"Other: {r[topic_other_col]}"
            else:
                topic_display = topic
            
            subject_counts[subj_display] += 1
            topics_by_subject[subj_display][topic_display] += 1
            
            comment_text = r[comment_col].strip()
            if comment_text:
                cat = categorize_comment(comment_text, categories)
                if cat is None:
                    kw1, kw2 = ("Uncategorized", "—")
                else:
                    kw1, kw2 = cat
                comments.append({
                    "subject": subj_display,
                    "topic": topic_display,
                    "comment": comment_text,
                    "kw1": kw1,
                    "kw2": kw2,
                })
        
        out[q.lower()] = {
            "subject_counts": subject_counts,
            "topics_by_subject": dict(topics_by_subject),
            "comments": comments,
        }
    
    # Q6, Q7, Q8: open-ended free text
    for q in ("Q6", "Q7", "Q8"):
        col = f"{q}_Open"
        comments = []
        for r in domain_rows:
            text = r[col].strip()
            if text:
                cat = categorize_comment(text, categories)
                if cat is None:
                    kw1, kw2 = ("Uncategorized", "—")
                else:
                    kw1, kw2 = cat
                comments.append({
                    "comment": text,
                    "kw1": kw1,
                    "kw2": kw2,
                })
        # Sort by kw1 then kw2 then comment
        comments.sort(key=lambda c: (c["kw1"], c["kw2"], c["comment"]))
        out[q.lower()] = comments
    
    return out


# ==================== CHART GENERATION ====================

def make_horizontal_bar_chart(counts, title, max_items=10):
    """Make a horizontal bar chart from a Counter, returns BytesIO PNG."""
    items = counts.most_common(max_items)
    if not items:
        return None
    
    labels = [item[0] for item in items]
    values = [item[1] for item in items]
    
    fig, ax = plt.subplots(figsize=(7, max(2.5, len(items) * 0.35)))
    y_pos = range(len(labels))
    bars = ax.barh(y_pos, values, color="#2E75B6")
    ax.set_yticks(y_pos)
    ax.set_yticklabels(labels, fontsize=9)
    ax.invert_yaxis()  # highest at top
    ax.set_xlabel("Number of Respondents")
    ax.set_title(title, fontsize=11, fontweight="bold")
    
    # Add value labels at end of bars
    for bar, val in zip(bars, values):
        ax.text(bar.get_width() + 0.1, bar.get_y() + bar.get_height()/2,
                str(val), va="center", fontsize=9)
    
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.tight_layout()
    
    buf = BytesIO()
    plt.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


# ==================== WORD DOC HELPERS ====================

def add_shading(cell, color_hex):
    """Add background color to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_section_header(doc, text, level=1):
    """Add a styled section header."""
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = COLOR_HEADER
        run.font.name = "Calibri"
    return p


def add_subject_header_row(doc, text, color_hex=COLOR_DARKER_BG):
    """Add a single-row colored header table for subject groupings."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    add_shading(cell, color_hex)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    return table


def add_two_column_section(doc, items, render_item):
    """
    Render `items` in a two-column table layout.
    `render_item(cell, item)` is called for each item to fill its cell.
    """
    if not items:
        return
    
    # Split into two columns
    midpoint = (len(items) + 1) // 2
    left_col = items[:midpoint]
    right_col = items[midpoint:]
    
    # Pad shorter column
    while len(right_col) < len(left_col):
        right_col.append(None)
    
    table = doc.add_table(rows=len(left_col), cols=2)
    table.autofit = False
    
    for row_idx, (left, right) in enumerate(zip(left_col, right_col)):
        row = table.rows[row_idx]
        if left is not None:
            render_item(row.cells[0], left)
        if right is not None:
            render_item(row.cells[1], right)


# ==================== REPORT GENERATION ====================

def render_comment_in_cell(cell, item):
    """Render a single categorized comment into a table cell."""
    cell.paragraphs[0].text = ""  # clear default
    p = cell.paragraphs[0]
    
    # Keyword tag line (bold, small)
    tag_run = p.add_run(f"[{item['kw1']} / {item['kw2']}]")
    tag_run.bold = True
    tag_run.font.size = Pt(8)
    tag_run.font.color.rgb = COLOR_ACCENT
    
    # Comment text
    p2 = cell.add_paragraph()
    body_run = p2.add_run(item["comment"])
    body_run.font.size = Pt(10)


def render_q3_q4_q5_comment_in_cell(cell, item):
    """Render a Q3/Q4/Q5 comment with subject/topic context."""
    cell.paragraphs[0].text = ""
    p = cell.paragraphs[0]
    
    tag_run = p.add_run(f"[{item['kw1']} / {item['kw2']}]")
    tag_run.bold = True
    tag_run.font.size = Pt(8)
    tag_run.font.color.rgb = COLOR_ACCENT
    
    p2 = cell.add_paragraph()
    ctx_run = p2.add_run(f"Topic: {item['topic']}")
    ctx_run.italic = True
    ctx_run.font.size = Pt(9)
    
    p3 = cell.add_paragraph()
    body_run = p3.add_run(item["comment"])
    body_run.font.size = Pt(10)


def add_chart_to_doc(doc, chart_buf, width_inches=6.5):
    """Insert a chart image into the document."""
    if chart_buf is None:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(chart_buf, width=Inches(width_inches))


def build_report(analysis, output_path):
    """Build a Word report for a single domain."""
    doc = Document()
    
    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    
    # ----- TITLE PAGE -----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("EVALUATION REPORT")
    title_run.font.size = Pt(28)
    title_run.bold = True
    title_run.font.color.rgb = COLOR_HEADER
    
    doc.add_paragraph()  # spacer
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(analysis["domain"])
    sub_run.font.size = Pt(20)
    sub_run.font.color.rgb = COLOR_ACCENT
    
    doc.add_paragraph()
    
    # Summary stats
    pct_first_time = (analysis["first_time"] / analysis["respondents"] * 100) if analysis["respondents"] else 0
    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sum_run = summary.add_run(
        f"{analysis['respondents']} respondents  •  "
        f"{analysis['first_time']} first-time attendees ({pct_first_time:.0f}%)"
    )
    sum_run.font.size = Pt(12)
    
    # Note about synthetic data
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note.add_run("[ DEMONSTRATION REPORT — SYNTHETIC DATA ONLY ]")
    note_run.italic = True
    note_run.font.size = Pt(10)
    note_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    
    doc.add_page_break()
    
    # ----- Q2: TOP SAFETY TOPICS -----
    add_section_header(doc, "Question 2: Top Safety Topics", level=1)
    doc.add_paragraph(
        "Respondents selected up to four safety topics they want to hear about "
        "at future meetings. Topics ranked by total selections."
    )
    chart = make_horizontal_bar_chart(
        analysis["q2_counts"],
        f"Top Safety Topics — {analysis['domain']}",
    )
    add_chart_to_doc(doc, chart)
    
    doc.add_page_break()
    
    # ----- Q3, Q4, Q5: SUBJECT/TOPIC + COMMENTS -----
    for q_num, q_label in [("q3", "Question 3"), ("q4", "Question 4"), ("q5", "Question 5")]:
        q_data = analysis[q_num]
        
        add_section_header(doc, f"{q_label}: Safety Concern", level=1)
        
        # Subject distribution chart
        if q_data["subject_counts"]:
            chart = make_horizontal_bar_chart(
                q_data["subject_counts"],
                f"{q_label} — Subject Distribution",
            )
            add_chart_to_doc(doc, chart, width_inches=5.5)
        
        # Topic breakdown - one chart per subject that has multiple topic responses
        if any(len(t) > 0 for t in q_data["topics_by_subject"].values()):
            doc.add_paragraph()
            sub_h = doc.add_paragraph()
            sub_h_run = sub_h.add_run("Topics chosen within each subject:")
            sub_h_run.bold = True
            sub_h_run.font.size = Pt(11)
            
            for subject, topic_counts in sorted(q_data["topics_by_subject"].items()):
                if not topic_counts:
                    continue
                add_subject_header_row(doc, subject)
                # Simple list of topic -> count
                tbl = doc.add_table(rows=len(topic_counts), cols=2)
                tbl.autofit = False
                for i, (topic, count) in enumerate(topic_counts.most_common()):
                    tbl.rows[i].cells[0].text = topic
                    tbl.rows[i].cells[1].text = str(count)
                    for cell in tbl.rows[i].cells:
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.font.size = Pt(10)
        
        # Further Comments (Qb)
        if q_data["comments"]:
            doc.add_paragraph()
            ch = doc.add_paragraph()
            ch_run = ch.add_run(f"{q_label.split()[1]}b. Further Comments")
            ch_run.bold = True
            ch_run.font.size = Pt(12)
            ch_run.font.color.rgb = COLOR_HEADER
            
            # Group by Keyword 1, then list within
            grouped = defaultdict(list)
            for c in q_data["comments"]:
                grouped[c["kw1"]].append(c)
            
            for kw1 in sorted(grouped.keys()):
                items = sorted(grouped[kw1], key=lambda x: (x["kw2"], x["comment"]))
                add_subject_header_row(doc, kw1, color_hex=COLOR_DARKER_BG)
                add_two_column_section(doc, items, render_q3_q4_q5_comment_in_cell)
                doc.add_paragraph()  # spacer
        
        doc.add_page_break()
    
    # ----- Q6, Q7, Q8: OPEN-ENDED -----
    open_questions = [
        ("q6", "Question 6: How can we evolve to make this more valuable?"),
        ("q7", "Question 7: Challenges or emerging risks over the next 3 years"),
        ("q8", "Question 8: General feedback"),
    ]
    
    for q_num, q_label in open_questions:
        comments = analysis[q_num]
        add_section_header(doc, q_label, level=1)
        
        if not comments:
            doc.add_paragraph("(No responses to this question.)")
            doc.add_page_break()
            continue
        
        # Group by Keyword 1
        grouped = defaultdict(list)
        for c in comments:
            grouped[c["kw1"]].append(c)
        
        for kw1 in sorted(grouped.keys()):
            items = sorted(grouped[kw1], key=lambda x: (x["kw2"], x["comment"]))
            add_subject_header_row(doc, kw1, color_hex=COLOR_DARKER_BG)
            add_two_column_section(doc, items, render_comment_in_cell)
            doc.add_paragraph()
        
        doc.add_page_break()
    
    # ----- APPENDIX: RAW STATS -----
    add_section_header(doc, "Appendix: Quick Statistics", level=1)
    doc.add_paragraph(f"Total respondents: {analysis['respondents']}")
    doc.add_paragraph(f"First-time attendees: {analysis['first_time']} ({pct_first_time:.0f}%)")
    doc.add_paragraph(f"Q2 total topic selections: {sum(analysis['q2_counts'].values())}")
    
    uncategorized_count = 0
    for q in ("q3", "q4", "q5"):
        for c in analysis[q]["comments"]:
            if c["kw1"] == "Uncategorized":
                uncategorized_count += 1
    for q in ("q6", "q7", "q8"):
        for c in analysis[q]:
            if c["kw1"] == "Uncategorized":
                uncategorized_count += 1
    
    if uncategorized_count > 0:
        p = doc.add_paragraph()
        run = p.add_run(
            f"\n⚠ {uncategorized_count} comment(s) could not be auto-categorized "
            "and need manual review. Search for [Uncategorized / —] tags in this document."
        )
        run.font.color.rgb = RGBColor(0xC0, 0x50, 0x4D)
        run.italic = True
    
    doc.save(output_path)


# ==================== MAIN ====================

def main():
    print("Loading data...")
    responses = load_responses(INPUT_CSV)
    categories = load_keywords(KEYWORDS_FILE)
    
    domains = sorted(set(r["Q1_Domain"] for r in responses))
    print(f"Found {len(responses)} responses across {len(domains)} domains.")
    
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    for domain in domains:
        print(f"\nProcessing: {domain}")
        analysis = analyze_domain(responses, domain, categories)
        
        # Sanitize domain name for filename
        safe_name = domain.replace(" ", "_").replace("-", "").replace("/", "_")
        safe_name = "".join(c for c in safe_name if c.isalnum() or c == "_")
        output_path = os.path.join(OUTPUT_DIR, f"{safe_name}_report.docx")
        
        build_report(analysis, output_path)
        print(f"  Wrote: {output_path}")
        print(f"  Respondents: {analysis['respondents']}")
        print(f"  First-time: {analysis['first_time']}")
    
    print(f"\n✓ Complete. Reports written to {OUTPUT_DIR}/")


if __name__ == "__main__":
    main()
